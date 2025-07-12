import os
import json
import docx
import csv
import pandas as pd
import pytesseract
import speech_recognition as sr
from PIL import Image
from agents.logger import get_logger

logger = get_logger("chat_agent", "logs/chat_agent.log")

class ChatAgent:
    def __init__(self, gpt_client):
        self.client = gpt_client
        self.selected_model_info = None
        self.request_alternative = False
        self.processing = False
        self.last_task_type = None
        self.chat_history = []

    def set_selected_model(self, model_dict):
        self.selected_model_info = model_dict

    def _read_text_file(self, file_path):
        try:
            with open(file_path, 'r') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""

    def _read_docx_file(self, file_path):
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return ""

    def _read_csv_file(self, file_path):
        try:
            with open(file_path, 'r', encoding="utf-8") as f:
                reader = csv.reader(f)
                return "\n".join([", ".join(row) for row in reader])
        except Exception as e:
            logger.error(f"Error reading CSV: {e}")
            return ""

    def _read_xlsx_file(self, file_path):
        try:
            df = pd.read_excel(file_path)
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Error reading XLSX: {e}")
            return ""

    def _read_json_file(self, file_path):
        try:
            with open(file_path, 'r') as f:
                return json.dumps(json.load(f), indent=2)
        except Exception as e:
            logger.error(f"Error reading JSON: {e}")
            return ""

    def _read_image_file(self, file_path):
        try:
            img = Image.open(file_path)
            return pytesseract.image_to_string(img)
        except Exception as e:
            logger.error(f"Error reading image: {e}")
            return ""

    def _read_audio_file(self, file_path):
        recognizer = sr.Recognizer()
        try:
            with sr.AudioFile(file_path) as source:
                audio_data = recognizer.record(source)
                return recognizer.recognize_google(audio_data)
        except Exception as e:
            logger.error(f"Error reading audio: {e}")
            return ""

    def _handle_file_input(self, file_path):
        ext = os.path.splitext(file_path)[-1].lower()
        if ext == '.txt':
            return self._read_text_file(file_path)
        elif ext == '.docx':
            return self._read_docx_file(file_path)
        elif ext == '.csv':
            return self._read_csv_file(file_path)
        elif ext == '.xlsx':
            return self._read_xlsx_file(file_path)
        elif ext == '.json':
            return self._read_json_file(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return self._read_image_file(file_path)
        elif ext in ['.mp3', '.wav']:
            return self._read_audio_file(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return ""

    def _extract_task_type(self, user_input):
        task_keywords = {
            "image_generation": ["image", "generate image", "picture", "visual", "dalle"],
            "text_generation": ["text", "generate text", "content", "paragraph"],
            "translation": ["translate", "translation", "language"],
            "summarization": ["summary", "summarize"],
            "audio_generation": ["speech", "voice", "generate audio", "tts"]
        }
        lowered = user_input.lower()
        for task, keywords in task_keywords.items():
            if any(k in lowered for k in keywords):
                return task
        return "general"

    def _append_to_history(self, role, content):
        self.chat_history.append({"role": role, "content": content})
        if len(self.chat_history) > 8:
            self.chat_history.pop(0)

    import os
import json
import docx
import csv
import pandas as pd
import pytesseract
import speech_recognition as sr
from PIL import Image
from agents.logger import get_logger

logger = get_logger("chat_agent", "logs/chat_agent.log")


class ChatAgent:
    def __init__(self, gpt_client):
        self.client = gpt_client
        self.selected_model_info = None
        self.request_alternative = False
        self.processing = False
        self.last_task_type = None
        self.chat_history = []

    def set_selected_model(self, model_dict):
        self.selected_model_info = model_dict

    def _read_text_file(self, file_path):
        try:
            with open(file_path, 'r') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return ""

    def _read_docx_file(self, file_path):
        try:
            doc = docx.Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            return ""

    def _read_csv_file(self, file_path):
        try:
            with open(file_path, ' encoding="utf-8') as f:
                reader = csv.reader(f)
                return "\n".join([", ".join(row) for row in reader])
        except Exception as e:
            logger.error(f"Error reading CSV: {e}")
            return ""

    def _read_xlsx_file(self, file_path):
        try:
            df = pd.read_excel(file_path)
            return df.to_string(index=False)
        except Exception as e:
            logger.error(f"Error reading XLSX: {e}")
            return ""

    def _read_json_file(self, file_path):
        try:
            with open(file_path, 'r') as f:
                return json.dumps(json.load(f), indent=2)
        except Exception as e:
            logger.error(f"Error reading JSON: {e}")
            return ""

    def _read_image_file(self, file_path):
        try:
            img = Image.open(file_path)
            return pytesseract.image_to_string(img)
        except Exception as e:
            logger.error(f"Error reading image: {e}")
            return ""

    def _read_audio_file(self, file_path):
        recognizer = sr.Recognizer()
        try:
            with sr.AudioFile(file_path) as source:
                audio_data = recognizer.record(source)
                return recognizer.recognize_google(audio_data)
        except Exception as e:
            logger.error(f"Error reading audio: {e}")
            return ""

    def _handle_file_input(self, file_path):
        ext = os.path.splitext(file_path)[-1].lower()
        if ext == '.txt':
            return self._read_text_file(file_path)
        elif ext == '.docx':
            return self._read_docx_file(file_path)
        elif ext == '.csv':
            return self._read_csv_file(file_path)
        elif ext == '.xlsx':
            return self._read_xlsx_file(file_path)
        elif ext == '.json':
            return self._read_json_file(file_path)
        elif ext in ['.png', '.jpg', '.jpeg']:
            return self._read_image_file(file_path)
        elif ext in ['.mp3', '.wav']:
            return self._read_audio_file(file_path)
        else:
            logger.warning(f"Unsupported file format: {ext}")
            return ""

    def _extract_task_type(self, user_input):
        task_keywords = {
            "image_generation": ["image", "generate image", "picture", "visual", "dalle"],
            "text_generation": ["text", "generate text", "content", "paragraph"],
            "translation": ["translate", "translation", "language"],
            "summarization": ["summary", "summarize"],
            "audio_generation": ["speech", "voice", "generate audio", "tts"]
        }
        lowered = user_input.lower()
        for task, keywords in task_keywords.items():
            if any(k in lowered for k in keywords):
                return task
        return "general"

    def _append_to_history(self, role, content):
        self.chat_history.append({"role": role, "content": content})
        if len(self.chat_history) > 8:
            self.chat_history.pop(0)

    def process_web_input(self, user_input):
        if self.processing:
            return {
                "proceed": False,
                "message": "Please wait, your previous request is still being processed.",
                "mode": "wait"
            }

        self.processing = True

        try:
            if not user_input or not user_input.strip():
                return {
                    "proceed": False,
                    "message": "Input is empty. Please provide a valid question or requirement.",
                    "mode": "empty"
                }

            lowered = user_input.lower().strip()

            if lowered in ["hi", "hello", "hey", "hola", "namaste"]:
                return {
                    "proceed": False,
                    "message": "Hello! I'm your AI assistant. Tell me what task you need help with—text, image, audio, summarization, or translation.",
                    "mode": "greeting"
                }

            if "other than this" in lowered or "another model" in lowered:
                self.request_alternative = True
                if not self.last_task_type:
                    return {
                        "proceed": False,
                        "message": "Please mention the task again (like 'generate image') so I can suggest another model.",
                        "mode": "missing_context"
                    }
                return {
                    "proceed": True,
                    "message": f"Okay! I’ll find an alternative model for your task: {self.last_task_type.replace('_', ' ')}.",
                    "mode": "alternative"
                }

            if any(keyword in lowered for keyword in ["price", "speed", "developer", "cloud", "accuracy", "region", "why"]):
                if self.selected_model_info:
                    answer = self.handle_follow_up(user_input)
                    self._append_to_history("user", user_input)
                    self._append_to_history("assistant", answer)
                    return {
                        "proceed": False,
                        "message": answer,
                        "mode": "follow_up"
                    }

            task_type = self._extract_task_type(user_input)
            self.last_task_type = task_type

            messages = [
                {
                    "role": "system",
                    "content": (
                        f"You are a helpful AI assistant that recommends the best AI model for user tasks like text generation, image generation, audio processing, translation, or summarization.\n"
                        f"If the user is asking about '{task_type.replace('_', ' ')}', suggest suitable AI models (e.g., GPT-4, Claude, DALL·E 3, Stable Diffusion, etc.) and explain why.\n"
                        f"Your goal is to explain your reasoning simply. Never repeat model names or prices unnecessarily. Be conversational and clear."
                    )
                }
            ]

            messages.extend(self.chat_history[-6:])
            messages.append({"role": "user", "content": user_input.strip()})

            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                temperature=0.4,
                max_tokens=350
            )

            result = response.choices[0].message.content.strip()
            logger.info("Web input analysis result:\n" + result)

            self._append_to_history("user", user_input.strip())
            self._append_to_history("assistant", result)

            return {
                "proceed": True,
                "message": result,
                "mode": "new_request"
            }

        except Exception as e:
            logger.error(f"Web GPT error: {repr(e)}")
            return {
                "proceed": False,
                "message": "⚠️ Sorry, something went wrong. Please try again.",
                "mode": "error"
            }

        finally:
            self.processing = False

    def handle_follow_up(self, user_query):
        if not self.selected_model_info:
            return "No model is currently selected. Please ask for a model recommendation first."

        model_json = json.dumps(self.selected_model_info, indent=2)

        messages = [
            {
                "role": "system",
                "content": (
                    f"You are an expert assistant explaining the chosen AI model below to the user. Use natural, friendly, helpful language.\n"
                    f"Model Details:\n{model_json}\n\n"
                    f"Give accurate, useful details about pricing, performance, speed, or region. If anything is missing, say 'not specified' but focus on what you do know."
                )
            },
            {"role": "user", "content": user_query}
        ]

        try:
            response = self.client.chat.completions.create(
                model="gpt-4o",
                messages=messages,
                temperature=0.4,
                max_tokens=300
            )
            return response.choices[0].message.content.strip()

        except Exception as e:
            logger.error(f"Follow-up GPT error: {e}")
            return "⚠️ Sorry, I couldn't answer that. Please try asking again."
